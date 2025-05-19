from typing import List
import logging
import re
from io import BytesIO
from langchain.schema import Document
import pypdf
import docx  # python-docx
import pyarabic.araby as araby
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


def load_pdf(file_path_or_bytes: str | bytes) -> str:
    """
    Loads text content from a PDF file.

    Args:
        file_path_or_bytes (str | bytes): Path to the PDF file or bytes content.

    Returns:
        str: The extracted text content as a single string.

    Raises:
        FileNotFoundError: If the file path is invalid.
        Exception: For errors during PDF parsing.
    """
    logger.debug(
        f"Loading PDF: {'from path' if isinstance(file_path_or_bytes, str) else 'from bytes'}"
    )
    text_content = ""
    try:
        if isinstance(file_path_or_bytes, str):
            reader = pypdf.PdfReader(file_path_or_bytes)
        else:
            reader = pypdf.PdfReader(BytesIO(file_path_or_bytes))

        for page in reader.pages:
            text_content += page.extract_text() + "\n"  # Add newline between pages
        text_content = normalize_arabic_text(text_content)
        logger.info(
            f"Successfully extracted text from PDF (pages: {len(reader.pages)})."
        )
        return text_content
    except FileNotFoundError:
        logger.error(f"PDF file not found: {file_path_or_bytes}")
        raise
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}", exc_info=True)
        raise Exception(f"Failed to parse PDF content: {e}")


def load_docx(file_path_or_bytes: str | bytes) -> str:
    """
    Loads text content from a DOCX file.

    Args:
        file_path_or_bytes (str | bytes): Path to the DOCX file or bytes content.

    Returns:
        str: The extracted text content.

    Raises:
        FileNotFoundError: If the file path is invalid.
        Exception: For errors during DOCX parsing.
    """
    logger.debug(
        f"Loading DOCX: {'from path' if isinstance(file_path_or_bytes, str) else 'from bytes'}"
    )
    try:
        document = docx.Document(file_path_or_bytes)
        text_content = "\n".join([para.text for para in document.paragraphs])
        logger.info(f"Successfully extracted text from DOCX.")
        return text_content
    except FileNotFoundError:
        logger.error(f"DOCX file not found: {file_path_or_bytes}")
        raise
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}", exc_info=True)
        raise Exception(f"Failed to parse DOCX content: {e}")


def load_wiki_page(url_or_content: str) -> str:
    """
    Loads text content from a Wiki page (placeholder - requires specific implementation).

    Args:
        url_or_content (str): URL or HTML/markup content.

    Returns:
        str: Extracted text content (or original content if not URL).
    """
    logger.warning(
        "load_wiki_page is a placeholder. Requires specific implementation (API/Scraping). Returning input."
    )
    # Placeholder: Implement actual Wiki loading (e.g., using requests and BeautifulSoup)
    # For now, just returns the input, assuming it might be pre-fetched content.
    return url_or_content

# In text_processing.py
# def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
#     ...

#     """
#     Extracts question-answer pairs from text and returns them as Documents with metadata.
    
#     Each Document contains the question text and metadata indicating it's a question, 
#     and links to its corresponding answer via the question_id.
#     """
#     pattern = r"(\d+\.\s*[^?\n]+?)\n([^1\d+.\n]+)(?=\n\d+\.|\Z)"
#     matches = re.finditer(pattern, text, re.DOTALL)

#     questions = []
#     answers_dict = {}

#     for match in matches:
#         q_num = len(questions) + 1
#         question = match.group(1).strip()
#         answer = match.group(2).strip()

#         questions.append(Document(
#             page_content=question,
#             metadata={
#                 "question_id": q_num,
#                 "type": "question"
#             }
#         ))
#         answers_dict[q_num] = answer

#     return questions

def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 150) -> List[str]:
    """
    Extracts Arabic Q&A pairs and returns them in clean, readable RTL format.

    Format:
    س: {question}
    ج: {answer}
    """
    import re

    if not text:
        return []

    def convert_to_arabic_digits(text):
        western_to_arabic = str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩")
        return text.translate(western_to_arabic)

    pattern = r"(?P<num>\d+)\.\s*(?P<q>.+?)\n(?P<a>.*?)(?=\n\d+\.|\Z)"
    matches = re.finditer(pattern, text, re.DOTALL)

    chunks = []
    for match in matches:
        num = convert_to_arabic_digits(match.group("num"))
        question = match.group("q").strip()
        answer = match.group("a").strip()
        rtl_q = f"\u200Fس: {num}. {question}"
        rtl_a = f"\u200Fج: {answer}"
        chunks.append(f"{rtl_q}\n{rtl_a}")

    if chunks:
        logger.info(f"Extracted {len(chunks)} Q&A chunks.")
        return chunks

    # fallback
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_text(text)
    logger.info(f"Split text into {len(chunks)} fallback chunks.")
    return chunks





def normalize_arabic_text(text: str) -> str:
    """
    Applies basic normalization techniques specific to Arabic text using pyarabic.

    Removes diacritics (tashkeel), normalizes alef variants, normalizes yaa variants, removes tatweel.

    Args:
        text (str): The input Arabic text.

    Returns:
        str: The normalized Arabic text.
    """
    if not text:
        return ""
    logger.debug("Normalizing Arabic text...")
    text = araby.strip_tashkeel(text)
    text = araby.strip_tatweel(text)
    # pyarabic normalize_alef replaces all variants with plain alef 'ا'
    text = araby.normalize_alef(text)
    # pyarabic normalize_hamza replaces different hamza forms (useful in some contexts)
    # text = araby.normalize_hamza(text)
    # Replace ending yaa variants with plain yaa 'ي'
    text = re.sub(r"ى$", "ي", text)  # Specific ending Alef Maksura to Yaa
    # Consider more specific rules if needed
    return text
